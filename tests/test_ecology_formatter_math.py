from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from format_ecology_submission_docx import (  # noqa: E402
    MATH_NS,
    _normalize_libreoffice_math_superscripts,
)


def test_optimum_star_is_normalized_without_touching_plain_asterisk() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("plain * stays plain")

    sup = OxmlElement("m:sup")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = "*"
    run.append(text)
    sup.append(run)
    paragraph._p.append(sup)

    replaced = _normalize_libreoffice_math_superscripts(doc)

    assert replaced == 1
    text_tag = f"{{{MATH_NS}}}t"
    assert "".join(node.text or "" for node in sup.iter(text_tag)) == "opt"
    assert paragraph.text == "plain * stays plain"


def test_non_star_math_superscript_is_unchanged() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()

    sup = OxmlElement("m:sup")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = "2"
    run.append(text)
    sup.append(run)
    paragraph._p.append(sup)

    assert _normalize_libreoffice_math_superscripts(doc) == 0
    text_tag = f"{{{MATH_NS}}}t"
    assert "".join(node.text or "" for node in sup.iter(text_tag)) == "2"
