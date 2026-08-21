from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TITLE_BREAK = "[[ECOLOGY_SECTION_BREAK_AFTER_TITLE]]"
REF_BREAK = "[[ECOLOGY_SECTION_BREAK_AFTER_REFERENCES]]"
PAGE_BREAK = "[[ECOLOGY_PAGE_BREAK]]"


def _remove_all_runs(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _suppress_line_numbers(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def _append_section_properties(paragraph, *, line_numbers: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:sectPr"))
    if old is not None:
        p_pr.remove(old)

    sect_pr = OxmlElement("w:sectPr")

    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "12240")
    pg_sz.set(qn("w:h"), "15840")
    sect_pr.append(pg_sz)

    pg_mar = OxmlElement("w:pgMar")
    for key, value in {
        "top": "1440",
        "right": "1440",
        "bottom": "1440",
        "left": "1440",
        "header": "720",
        "footer": "720",
        "gutter": "0",
    }.items():
        pg_mar.set(qn(f"w:{key}"), value)
    sect_pr.append(pg_mar)

    if line_numbers:
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:start"), "1")
        ln.set(qn("w:restart"), "continuous")
        ln.set(qn("w:distance"), "360")
        sect_pr.append(ln)

    p_pr.append(sect_pr)
    _remove_all_runs(paragraph)


def _replace_pagebreak_marker(paragraph) -> None:
    _remove_all_runs(paragraph)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _set_run_font(run, size: Pt) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")


def _set_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    _set_run_font(run, Pt(10))


def _configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    for i, section in enumerate(doc.sections):
        if i > 0:
            section.footer.is_linked_to_previous = True
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _remove_all_runs(p)
    _set_page_field(p)


def _format_document(doc: Document, *, appendix: bool) -> None:
    region = "appendix" if appendix else "title"

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not appendix and text == TITLE_BREAK:
            _suppress_line_numbers(paragraph)
            _append_section_properties(paragraph, line_numbers=False)
            region = "numbered"
            continue

        if not appendix and text == REF_BREAK:
            _append_section_properties(paragraph, line_numbers=True)
            _suppress_line_numbers(paragraph)
            region = "postrefs"
            continue

        if not appendix and region in {"title", "postrefs"}:
            # LibreOffice can apply section line-number settings more broadly than
            # Word. Explicit paragraph suppression guarantees that the rendered
            # review file shows numbers only from Abstract through References.
            _suppress_line_numbers(paragraph)

        if text == PAGE_BREAK:
            _replace_pagebreak_marker(paragraph)

    # All tables are post-References in the Ecology Main Document. Suppress line
    # numbers explicitly within table cells as a renderer-independent safeguard.
    if not appendix:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _suppress_line_numbers(paragraph)

    _configure_sections(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 2

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 2
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, Pt(12))

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        _set_run_font(run, Pt(10))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--appendix", action="store_true")
    args = parser.parse_args()

    doc = Document(args.input_docx)
    _format_document(doc, appendix=args.appendix)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)


if __name__ == "__main__":
    main()
